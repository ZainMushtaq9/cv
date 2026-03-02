import os
import io
import concurrent.futures
from flask import Flask, render_template, request, jsonify, send_file
from flask_cors import CORS
from werkzeug.utils import secure_filename
from dotenv import load_dotenv

from extractor import extract_text
from ai_engine import analyze_resume
from scoring import normalize_score
from ranking import rank_candidates

load_dotenv()

app = Flask(__name__)
CORS(app)

app.config['SECRET_KEY'] = os.environ.get('SECRET_KEY', 'super-secret-key')
app.config['GROQ_API_KEY'] = os.environ.get('GROQ_API_KEY')
app.config['MAX_CONTENT_LENGTH'] = 5 * 1024 * 1024  # 5 MB

UPLOAD_FOLDER = os.path.join(os.getcwd(), 'uploads')
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

ALLOWED_EXTENSIONS = {'txt', 'pdf', 'png', 'jpg', 'jpeg', 'docx'}

# Industry-specific keyword suggestions
INDUSTRY_SKILLS = {
    'technology': ['Python', 'JavaScript', 'SQL', 'AWS', 'Docker', 'Git', 'React', 'Node.js', 'API', 'Agile', 'CI/CD', 'Kubernetes', 'Cloud Computing', 'Machine Learning'],
    'finance': ['Financial Analysis', 'Excel', 'Bloomberg', 'Risk Management', 'Accounting', 'GAAP', 'Forecasting', 'Budgeting', 'SAP', 'Auditing', 'Compliance', 'Portfolio Management'],
    'healthcare': ['HIPAA', 'EMR', 'Patient Care', 'Clinical Research', 'Medical Terminology', 'Compliance', 'Epic', 'ICD-10', 'CPR', 'Healthcare Management', 'Pharmacy', 'Nursing'],
    'marketing': ['SEO', 'Google Analytics', 'Social Media', 'Content Strategy', 'CRM', 'HubSpot', 'Email Marketing', 'Branding', 'PPC', 'Copywriting', 'Market Research', 'Adobe Creative Suite'],
    'education': ['Curriculum Development', 'Classroom Management', 'Student Assessment', 'Lesson Planning', 'EdTech', 'Differentiated Instruction', 'IEP', 'Special Education', 'Teaching'],
    'engineering': ['CAD', 'AutoCAD', 'SolidWorks', 'Project Management', 'Quality Assurance', 'Lean Manufacturing', 'Six Sigma', 'MATLAB', 'Technical Drawing', 'Safety Compliance'],
    'data_science': ['Python', 'R', 'SQL', 'Machine Learning', 'TensorFlow', 'Pandas', 'Data Visualization', 'Statistical Analysis', 'Deep Learning', 'NLP', 'Tableau', 'Big Data'],
    'hr': ['Recruitment', 'Onboarding', 'Employee Relations', 'HRIS', 'Payroll', 'Performance Management', 'Compliance', 'Training & Development', 'Labor Law', 'Benefits Administration'],
    'sales': ['CRM', 'Salesforce', 'Lead Generation', 'Negotiation', 'Cold Calling', 'Account Management', 'Pipeline Management', 'B2B', 'Revenue Growth', 'Client Relations'],
    'design': ['Figma', 'Adobe Photoshop', 'Illustrator', 'UI/UX', 'Wireframing', 'Prototyping', 'Typography', 'Color Theory', 'User Research', 'Responsive Design'],
    'web_development': ['HTML', 'CSS', 'JavaScript', 'React', 'Node.js', 'TypeScript', 'REST API', 'MongoDB', 'PostgreSQL', 'Git', 'Webpack', 'Tailwind CSS'],
    'administrative': ['Microsoft Office', 'Data Entry', 'Scheduling', 'Filing', 'Communication', 'Organizational Skills', 'Time Management', 'Record Keeping', 'Customer Service'],
}


def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS


def save_and_extract(file, folder):
    if not file or not allowed_file(file.filename):
        return None
    filename = secure_filename(file.filename)
    filepath = os.path.join(folder, filename)
    file.save(filepath)
    text = extract_text(filepath)
    try:
        os.remove(filepath)
    except:
        pass
    return text


# ============ PAGES ============

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/individual')
def individual():
    return render_template('individual.html')

@app.route('/bulk')
def bulk():
    return render_template('bulk.html')

@app.route('/compare')
def compare():
    return render_template('compare.html')

@app.route('/about')
def about():
    return render_template('about.html')

@app.route('/contact')
def contact():
    return render_template('contact.html')

@app.route('/privacy')
def privacy():
    return render_template('privacy.html')

@app.route('/terms')
def terms():
    return render_template('terms.html')


# ============ API ENDPOINTS ============

@app.route('/api/analyze/single', methods=['POST'])
def analyze_single():
    api_key = app.config['GROQ_API_KEY']
    resume_file = request.files.get('resume')
    jd_file = request.files.get('jd_file')
    jd_text = request.form.get('jd_text', '')
    resume_text_input = request.form.get('resume_text', '')
    category = request.form.get('category', '')

    # Get resume content — from file upload OR pasted text
    resume_content = None
    if resume_file and resume_file.filename:
        resume_content = save_and_extract(resume_file, UPLOAD_FOLDER)
    if not resume_content and resume_text_input.strip():
        resume_content = resume_text_input.strip()
    if not resume_content:
        return jsonify({"status": "error", "message": "Resume is required. Upload a file or paste your resume text."}), 400

    # Get JD content — from file upload OR pasted text
    final_jd_text = jd_text
    if jd_file and jd_file.filename and allowed_file(jd_file.filename):
        extracted_jd = save_and_extract(jd_file, UPLOAD_FOLDER)
        if extracted_jd:
            final_jd_text = (extracted_jd + "\n" + final_jd_text).strip()
        else:
            if not final_jd_text.strip():
                return jsonify({"status": "error", "message": "Failed to extract text from the uploaded Job Description file. Please try pasting the text instead."}), 400

    if not final_jd_text.strip():
        return jsonify({"status": "error", "message": "Job Description is required. Please paste text or upload a file."}), 400

    # Get industry-specific skills
    category_key = category.lower().replace(' ', '_').replace('/', '_')
    category_skills = INDUSTRY_SKILLS.get(category_key, [])

    result = analyze_resume(api_key, resume_content, final_jd_text, category_skills)
    if not result:
        return jsonify({"status": "error", "message": "Failed to analyze resume. Please try again later."}), 500

    result = normalize_score(result)

    # Include resume text for frontend highlighting
    result['resume_text'] = resume_content

    return jsonify({"status": "success", "data": result})


@app.route('/api/analyze/bulk', methods=['POST'])
def analyze_bulk_api():
    api_key = app.config['GROQ_API_KEY']
    resume_files = request.files.getlist('resumes')
    jd_file = request.files.get('jd_file')
    jd_text = request.form.get('jd_text', '')

    if not resume_files or len(resume_files) == 0:
        return jsonify({"status": "error", "message": "Resume files are required"}), 400
    if len(resume_files) > 50:
        return jsonify({"status": "error", "message": "Maximum 50 resumes allowed per batch"}), 400

    final_jd_text = jd_text
    if jd_file and jd_file.filename and allowed_file(jd_file.filename):
        extracted_jd = save_and_extract(jd_file, UPLOAD_FOLDER)
        if extracted_jd:
            final_jd_text = (extracted_jd + "\n" + final_jd_text).strip()
        else:
            if not final_jd_text.strip():
                return jsonify({"status": "error", "message": "Failed to extract text from the uploaded Job Description file."}), 400

    if not final_jd_text.strip():
        return jsonify({"status": "error", "message": "Job Description is required."}), 400

    resumes_data = []
    for file in resume_files:
        if file and allowed_file(file.filename):
            filename = secure_filename(file.filename)
            filepath = os.path.join(UPLOAD_FOLDER, filename)
            file.save(filepath)
            text = extract_text(filepath)
            try:
                os.remove(filepath)
            except:
                pass
            if text:
                resumes_data.append({"filename": filename, "text": text})

    def process_candidate(candidate_data):
        res = analyze_resume(api_key, candidate_data['text'], final_jd_text)
        if res:
            res = normalize_score(res)
            res['filename'] = candidate_data['filename']
        return res

    results = []
    with concurrent.futures.ThreadPoolExecutor(max_workers=5) as executor:
        future_to_candidate = {executor.submit(process_candidate, c): c for c in resumes_data}
        for future in concurrent.futures.as_completed(future_to_candidate):
            try:
                candidate_result = future.result()
                if candidate_result:
                    results.append(candidate_result)
            except Exception:
                pass

    ranked_results = rank_candidates(results)
    return jsonify({"status": "success", "data": ranked_results})


@app.route('/api/analyze/compare', methods=['POST'])
def analyze_compare():
    """Compare one resume against multiple job descriptions."""
    api_key = app.config['GROQ_API_KEY']
    resume_file = request.files.get('resume')
    resume_text_input = request.form.get('resume_text', '')

    # Get resume content
    resume_content = None
    if resume_file and resume_file.filename:
        resume_content = save_and_extract(resume_file, UPLOAD_FOLDER)
    if not resume_content and resume_text_input.strip():
        resume_content = resume_text_input.strip()
    if not resume_content:
        return jsonify({"status": "error", "message": "Resume is required."}), 400

    # Collect multiple JDs from form
    jd_texts = []
    jd_labels = []
    for i in range(1, 6):  # Support up to 5 JDs
        jd = request.form.get(f'jd_text_{i}', '').strip()
        label = request.form.get(f'jd_label_{i}', f'Job {i}').strip()
        if jd:
            jd_texts.append(jd)
            jd_labels.append(label)

    if len(jd_texts) < 2:
        return jsonify({"status": "error", "message": "Please provide at least 2 job descriptions to compare."}), 400

    def process_jd(idx):
        res = analyze_resume(api_key, resume_content, jd_texts[idx])
        if res:
            res = normalize_score(res)
            res['jd_label'] = jd_labels[idx]
            res['jd_index'] = idx + 1
        return res

    results = []
    with concurrent.futures.ThreadPoolExecutor(max_workers=5) as executor:
        future_map = {executor.submit(process_jd, i): i for i in range(len(jd_texts))}
        for future in concurrent.futures.as_completed(future_map):
            try:
                r = future.result()
                if r:
                    results.append(r)
            except Exception:
                pass

    # Sort by match_score descending
    results.sort(key=lambda x: float(x.get('match_score', 0)), reverse=True)
    for idx, r in enumerate(results):
        r['rank'] = idx + 1

    return jsonify({"status": "success", "data": results})


@app.route('/api/industry/benchmarks', methods=['GET'])
def industry_benchmarks():
    """Return industry benchmark keywords."""
    industry = request.args.get('industry', '').lower().replace(' ', '_').replace('/', '_')
    keywords = INDUSTRY_SKILLS.get(industry, [])
    return jsonify({
        "status": "success",
        "industry": industry,
        "benchmark_keywords": keywords,
        "total": len(keywords)
    })


@app.route('/api/export/pdf', methods=['POST'])
def export_pdf():
    """Generate a PDF report from the analysis results."""
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib import colors
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
        from reportlab.lib.units import inch

        data = request.get_json()
        if not data:
            return jsonify({"status": "error", "message": "No data provided"}), 400

        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4, topMargin=0.5*inch, bottomMargin=0.5*inch)
        styles = getSampleStyleSheet()
        story = []

        # Title
        title_style = ParagraphStyle('CustomTitle', parent=styles['Title'], fontSize=18, textColor=colors.HexColor('#4338ca'))
        story.append(Paragraph("Resume Keyword Optimizer — Analysis Report", title_style))
        story.append(Spacer(1, 12))

        # Candidate & Score
        name = data.get('candidate_name', 'Unknown')
        score = data.get('match_score', 0)
        decision = data.get('decision', 'N/A')
        story.append(Paragraph(f"<b>Candidate:</b> {name}", styles['Normal']))
        story.append(Paragraph(f"<b>Match Score:</b> {score}% — <b>{decision}</b>", styles['Normal']))
        story.append(Spacer(1, 12))

        def add_section(title, items, color):
            h_style = ParagraphStyle('SectionH', parent=styles['Heading2'], textColor=colors.HexColor(color))
            story.append(Paragraph(title, h_style))
            if items:
                for item in items:
                    if isinstance(item, dict):
                        keyword = item.get('keyword', '')
                        detail = item.get('suggestion', item.get('reason', ''))
                        story.append(Paragraph(f"• <b>{keyword}</b>: {detail}", styles['Normal']))
                    else:
                        story.append(Paragraph(f"• {item}", styles['Normal']))
            else:
                story.append(Paragraph("None identified.", styles['Normal']))
            story.append(Spacer(1, 8))

        add_section("Matched Keywords", data.get('matched_keywords', []), '#16a34a')
        add_section("Missing Keywords", data.get('missing_keywords', []), '#dc2626')
        add_section("Extra Keywords (Resume-only)", data.get('extra_keywords', []), '#2563eb')
        add_section("How to Add Missing Keywords", data.get('missing_keyword_suggestions', []), '#9333ea')
        add_section("Strengths", data.get('strengths', []), '#16a34a')
        add_section("Weaknesses", data.get('weaknesses', []), '#dc2626')
        add_section("CV Improvement Tips", data.get('cv_improvement_suggestions', []), '#2563eb')

        # Reasoning
        story.append(Paragraph("<b>HR Reasoning:</b>", styles['Heading2']))
        story.append(Paragraph(data.get('reasoning', 'No reasoning provided.'), styles['Normal']))

        doc.build(story)
        buffer.seek(0)

        return send_file(buffer, as_attachment=True, download_name='resume_keyword_report.pdf', mimetype='application/pdf')

    except ImportError:
        return jsonify({"status": "error", "message": "PDF export requires 'reportlab'. Install with: pip install reportlab"}), 500
    except Exception as e:
        return jsonify({"status": "error", "message": f"PDF generation failed: {str(e)}"}), 500


@app.route('/api/export/word', methods=['POST'])
def export_word():
    """Generate a Word document report from the analysis results."""
    try:
        from docx import Document as DocxDocument
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        data = request.get_json()
        if not data:
            return jsonify({"status": "error", "message": "No data provided"}), 400

        doc = DocxDocument()

        title = doc.add_heading('Resume Keyword Optimizer — Analysis Report', level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph(f"Candidate: {data.get('candidate_name', 'Unknown')}")
        doc.add_paragraph(f"Match Score: {data.get('match_score', 0)}% — {data.get('decision', 'N/A')}")
        doc.add_paragraph('')

        def add_keyword_section(heading, items, color_hex):
            h = doc.add_heading(heading, level=2)
            if h.runs:
                h.runs[0].font.color.rgb = RGBColor.from_string(color_hex[1:])
            if items:
                for item in items:
                    if isinstance(item, dict):
                        p = doc.add_paragraph(style='List Bullet')
                        run = p.add_run(f"{item.get('keyword', '')}: ")
                        run.bold = True
                        p.add_run(item.get('suggestion', item.get('reason', '')))
                    else:
                        doc.add_paragraph(str(item), style='List Bullet')
            else:
                doc.add_paragraph("None identified.")

        add_keyword_section("Matched Keywords", data.get('matched_keywords', []), '#16a34a')
        add_keyword_section("Missing Keywords", data.get('missing_keywords', []), '#dc2626')
        add_keyword_section("Extra Keywords (Resume-only)", data.get('extra_keywords', []), '#2563eb')
        add_keyword_section("How to Add Missing Keywords", data.get('missing_keyword_suggestions', []), '#9333ea')
        add_keyword_section("Strengths", data.get('strengths', []), '#16a34a')
        add_keyword_section("Weaknesses", data.get('weaknesses', []), '#dc2626')
        add_keyword_section("CV Improvement Tips", data.get('cv_improvement_suggestions', []), '#2563eb')

        doc.add_heading("HR Reasoning", level=2)
        doc.add_paragraph(data.get('reasoning', 'No reasoning provided.'))

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        return send_file(buffer, as_attachment=True, download_name='resume_keyword_report.docx',
                         mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

    except ImportError:
        return jsonify({"status": "error", "message": "Word export requires 'python-docx'. Install with: pip install python-docx"}), 500
    except Exception as e:
        return jsonify({"status": "error", "message": f"Word generation failed: {str(e)}"}), 500


if __name__ == '__main__':
    app.run(debug=True, port=5001, use_reloader=False)
