import os
import logging
import pdfplumber
import docx
from PIL import Image

try:
    import cv2
    CV2_AVAILABLE = True
except (ImportError, AttributeError):
    CV2_AVAILABLE = False

try:
    import pytesseract
    TESSERACT_AVAILABLE = True
except ImportError:
    TESSERACT_AVAILABLE = False

try:
    import pypdfium2 as pdfium
    PDFIUM_AVAILABLE = True
except ImportError:
    PDFIUM_AVAILABLE = False

logger = logging.getLogger(__name__)


def extract_text(filepath):
    if not os.path.exists(filepath):
        return ""
    ext = filepath.rsplit('.', 1)[-1].lower()
    try:
        if ext == 'pdf':
            return extract_from_pdf(filepath)
        elif ext == 'docx':
            return extract_from_docx(filepath)
        elif ext == 'txt':
            return extract_from_txt(filepath)
        elif ext in ['png', 'jpg', 'jpeg']:
            return extract_from_image(filepath)
        else:
            logger.error(f"Unsupported file extension: {ext}")
            return ""
    except Exception as e:
        logger.error(f"Error extracting text from {filepath}: {str(e)}")
        return ""


def extract_from_pdf(filepath):
    text = ""
    try:
        with pdfplumber.open(filepath) as pdf:
            for page in pdf.pages:
                extr = page.extract_text()
                if extr:
                    text += extr + "\n"
    except Exception as e:
        logger.warning(f"pdfplumber extraction failed: {e}")

    if text.strip() and len(text.strip()) > 20:
        return text.strip()

    logger.info(f"No text found in PDF, attempting OCR on: {filepath}")
    return ocr_pdf(filepath)


def ocr_pdf(filepath):
    if not TESSERACT_AVAILABLE:
        logger.error("pytesseract not available for OCR")
        return ""
    ocr_text = ""
    if PDFIUM_AVAILABLE:
        try:
            pdf = pdfium.PdfDocument(filepath)
            for i in range(len(pdf)):
                page = pdf[i]
                bitmap = page.render(scale=300 / 72)
                pil_image = bitmap.to_pil()
                page_text = pytesseract.image_to_string(pil_image)
                if page_text:
                    ocr_text += page_text + "\n"
            pdf.close()
            if ocr_text.strip():
                return ocr_text.strip()
        except Exception as e:
            logger.warning(f"pypdfium2 OCR failed: {e}")
    try:
        from pdf2image import convert_from_path
        images = convert_from_path(filepath, dpi=300)
        for img in images:
            page_text = pytesseract.image_to_string(img)
            if page_text:
                ocr_text += page_text + "\n"
        if ocr_text.strip():
            return ocr_text.strip()
    except ImportError:
        logger.warning("pdf2image not available")
    except Exception as e:
        logger.warning(f"pdf2image OCR failed: {e}")
    return ocr_text.strip()


def extract_from_docx(filepath):
    doc = docx.Document(filepath)
    text = "\n".join([para.text for para in doc.paragraphs])
    return text.strip()


def extract_from_txt(filepath):
    with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
        return f.read().strip()


def extract_from_image(filepath):
    if not TESSERACT_AVAILABLE:
        logger.error("pytesseract not available for image OCR")
        return ""
    try:
        if CV2_AVAILABLE:
            img = cv2.imread(filepath)
            if img is not None:
                gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
                thresh = cv2.adaptiveThreshold(
                    gray, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C,
                    cv2.THRESH_BINARY, 11, 2
                )
                text = pytesseract.image_to_string(thresh)
                if text.strip():
                    return text.strip()
        img_pil = Image.open(filepath)
        if img_pil.mode != 'RGB':
            img_pil = img_pil.convert('RGB')
        return pytesseract.image_to_string(img_pil).strip()
    except Exception as e:
        logger.error(f"Image OCR error: {str(e)}")
        return ""
